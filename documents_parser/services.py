from collections import defaultdict
from datetime import datetime
import tempfile
import zipfile

import pytz

from documents_parser.serializers import ExtractedDataSerializer
from .models import ExtractedData  # Import your Django model
from .excel_controller import *
from .pdf_controller import *
from .github_controller import *
import uuid

from docx import Document
from io import BytesIO
from django.utils.html import escape
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX  
from docx.shared import Inches


def process_uploaded_file(file_path, uploaded_file_name, tag_names, user):
    results_by_tag_exact = {tag: [] for tag in tag_names}
    results_by_tag_partial = {tag: [] for tag in tag_names}
    file_type=None
    if uploaded_file_name.endswith('.pdf'):
        file_type="PDF"
        # Process PDF files
        text_data = read_pdf_file(file_path)
        for tag in tag_names:
            result_exact, result_partial = search_pdf(text_data, tag, uploaded_file_name, "PDF", user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.docx'):
        file_type="Word"
        text_data = read_doc_file(file_path)
        for tag in tag_names:
            result_exact, result_partial = search_pdf(text_data, tag, uploaded_file_name, "Word", user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.xls') or uploaded_file_name.endswith('.xlsx'):
        file_type="Excel"
        for tag in tag_names:
            result_exact, result_partial = extract_text_from_excel(file_path, tag, uploaded_file_name, user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.zip'):


        with tempfile.TemporaryDirectory() as extract_dir:
            # Extract the ZIP file
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)  # Extract all files to the temp directory
            
            # Set file type and prepare for processing
            file_type = "Github"
            zip_file_name = uploaded_file_name.split(".")[0]

            # Now read all files in the extracted directory
            data_dict = read_all_files_in_directory(extract_dir)  # Use the extracted directory
            # Search through the extracted files for the tags
            for tag in tag_names:
                result_exact, result_partial = search_github(tag, zip_file_name, data_dict)
                results_by_tag_exact[tag] += result_exact
                results_by_tag_partial[tag] += result_partial

    else:
        raise ValueError(f"Unsupported file format: {uploaded_file_name}")

    return file_type, results_by_tag_exact, results_by_tag_partial



def save_results_to_db(results,results_partial, file_name, file_type, user):

    search_id = uuid.uuid4()


    for tag, result_list in results.items():
        for result in result_list:
            saved_result = ExtractedData.objects.create(
                search_id=search_id,  # Assign the generated search ID
                source_file_name=result["Source File Name"],
                file_type=result["File Type"],
                tag_searched=tag,
                block_record=result['Block/Record'],
                location_of_tag=result['Location of the Tag'],
                # date_of_search is automatically set with auto_now_add=True
                search_author=user,
                match_type="exact",
                other=result.get('Other', ''),
                # line_id will be automatically generated in the save method if not provided
            )
            # Append the saved instance to the list
            # saved_results_exact.append(saved_result)
    
    
    for tag, result_list in results_partial.items():
        for result in result_list:
            saved_result = ExtractedData.objects.create(
                search_id=search_id,  # Assign the generated search ID
                source_file_name=result["Source File Name"],
                file_type=result["File Type"],
                tag_searched=tag,
                block_record=result['Block/Record'],
                location_of_tag=result['Location of the Tag'],
                # date_of_search is automatically set with auto_now_add=True
                search_author=user,
                match_type="partial",
                other=result.get('Other', ''),
                # line_id will be automatically generated in the save method if not provided
            )
            # Append the saved instance to the list
            # saved_results_partial.append(saved_result)


    return search_id



def find_records(source_file_name=None, file_type=None, tag_searched=None, start_date=None, end_date=None, search_author=None):

    records = ExtractedData.objects.all()

    # Apply filters based on provided arguments
    if source_file_name:
        records = records.filter(source_file_name__icontains=source_file_name)

    if file_type:
        records = records.filter(file_type__iexact=file_type)

    if tag_searched:
        records = records.filter(tag_searched__icontains=tag_searched)

    if start_date:
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date()
        records = records.filter(date_of_search__gte=start_date_obj)

    if end_date:
        end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date()
        records = records.filter(date_of_search__lte=end_date_obj)

    if search_author:
        records = records.filter(search_author__icontains=search_author)
    
    formatted_results = defaultdict(lambda: {'exact_matches': [], 'partial_matches': []})

    # Iterate through the results and group them
    for result in records:
        # Determine the match type and append the result to the correct category
        if result.match_type == 'exact':
            formatted_results[result.source_file_name]['exact_matches'].append(result)
        elif result.match_type == 'partial':
            formatted_results[result.source_file_name]['partial_matches'].append(result)
    return records


def export_search_results_to_word(serialized_formatted_results, search_id, search_date, user_name):
    document = Document()
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    # Add title
    document_title = f'Search Results for Search ID {search_id} on {search_date.strftime("%B %d, %Y %H:%M")} by {user_name}'
    document.add_heading(document_title, level=1)

    # List files
    document.add_heading('Files in Search', level=2)
    for file_name in serialized_formatted_results.keys():
        document.add_paragraph(file_name, style='List Bullet')

    # Define headers for tables
    headers = ['Source File Name', 'File Type', 'Tag Searched', 'Block/Record', 'Location of the Tag', 'Date of Search', 'Author Name', 'Other']

    # Process each file's results
    for file_name, matches in serialized_formatted_results.items():
        document.add_heading(f'Exact Matches for {file_name}', level=2)
        table_exact = document.add_table(rows=1, cols=len(headers))
        table_exact.style = 'Table Grid'
        hdr_cells_exact = table_exact.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells_exact[i].text = header

        # Fill the Exact Matches table
        for result in matches['exact_matches']:
            row_cells = table_exact.add_row().cells
            row_cells[0].text = clean_text(result['source_file_name'])
            row_cells[1].text = clean_text(result['file_type'])
            row_cells[2].text = clean_text(result['tag_searched'])
            row_cells[3].text = clean_text(result['block_record'])
            row_cells[4].text = clean_text(result['location_of_tag'])
            row_cells[5].text = datetime.strptime(result['date_of_search'],"%Y-%m-%dT%H:%M:%S.%fZ").strftime("%B %d, %Y %H:%M")
            row_cells[6].text = clean_text(user_name)
            row_cells[7].text = clean_text(result['other'])
            # Highlight the tag in the 'Block/Record' cell
            highlight_text(row_cells[3].paragraphs[0], result['tag_searched'],WD_COLOR_INDEX.YELLOW)

        document.add_heading(f'Partial Matches for {file_name}', level=2)
        table_partial = document.add_table(rows=1, cols=len(headers))
        table_partial.style = 'Table Grid'
        hdr_cells_partial = table_partial.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells_partial[i].text = header

        # Fill the Partial Matches table
        for result in matches['partial_matches']:
            row_cells = table_partial.add_row().cells
            row_cells[0].text = clean_text(result['source_file_name'])
            row_cells[1].text = clean_text(result['file_type'])
            row_cells[2].text = clean_text(result['tag_searched'])
            row_cells[3].text = clean_text(result['block_record'])
            row_cells[4].text = clean_text(result['location_of_tag'])
            row_cells[5].text = datetime.strptime(result['date_of_search'],"%Y-%m-%dT%H:%M:%S.%fZ").strftime("%B %d, %Y %H:%M")
            row_cells[6].text = clean_text(user_name)
            row_cells[7].text = clean_text(result['other'])
            
            # Highlight the tag in the 'Block/Record' cell
            highlight_text(row_cells[3].paragraphs[0], result['tag_searched'],WD_COLOR_INDEX.GRAY_50)

    # Calculate summary statistics using the new function
    summary_statistics = calculate_summary_statistics(serialized_formatted_results)

    # Add summary table
    document.add_heading('Summary Statistics', level=2)
    summary_table = document.add_table(rows=1, cols=len(headers) + 1)
    summary_table.style = 'Table Grid'

    # Prepare header row for summary table
    summary_hdr_cells = summary_table.rows[0].cells
    summary_hdr_cells[0].text = 'File Name / Tag'
    tag_list = list({tag for tags in summary_statistics.values() for tag in tags.keys()})
    for i, tag in enumerate(tag_list):
        summary_hdr_cells[i + 1].text = tag

    # Fill summary table with tag counts
    for file_name, tags in summary_statistics.items():
        row_cells = summary_table.add_row().cells
        row_cells[0].text = file_name
        for i, tag in enumerate(tag_list):
            tag_data = tags.get(tag, {'exact': 0, 'partial': 0})
            row_cells[i + 1].text = f"Exact: {tag_data['exact']} / Partial: {tag_data['partial']}"

    # Save document to a BytesIO object
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def highlight_text(paragraph, tag, color=WD_COLOR_INDEX.YELLOW):
    text=paragraph.text.lower()
    parts = text.split(tag.lower())


    if len(parts) <= 1:
        return  # Tag not found, nothing to highlight

    # Clear existing paragraph text
    paragraph.clear()

    # Reconstruct paragraph with highlighted tag
    for i, part in enumerate(parts):
        if i > 0:
            # Create a run for the highlighted tag
            run = paragraph.add_run(tag)
            run.font.highlight_color =  color # Apply text highlight color using WD_COLOR_INDEX

        # Add the normal text
        paragraph.add_run(part)





def append_dicts(dict1, dict2):
    # Ensure both dictionaries have the same keys
    if dict1.keys() != dict2.keys():
        raise ValueError("Both dictionaries must have the same keys.")

    # Create a new dictionary to store the appended values
    result_dict = {}

    for key in dict1.keys():
        # Check if the value is a list; if not, convert to list for appending
        value1 = dict1[key] if isinstance(dict1[key], list) else [dict1[key]]
        value2 = dict2[key] if isinstance(dict2[key], list) else [dict2[key]]

        # Append values from dict2 to dict1 under the same key
        result_dict[key] = value1 + value2
    
    

    return result_dict


def calculate_summary_statistics(serialized_formatted_results):

    summary_statistics = defaultdict(lambda: defaultdict(lambda: {'exact': 0, 'partial': 0}))

    # Calculate tag counts for exact and partial matches
    for file_name, matches in serialized_formatted_results.items():
        for result in matches['exact_matches']:
            summary_statistics[file_name][result['tag_searched']]['exact'] += 1
        for result in matches['partial_matches']:
            summary_statistics[file_name][result['tag_searched']]['partial'] += 1

    # Convert to regular dictionary for easier manipulation and output
    formatted_summary_statistics = {file: dict(tag_counts) for file, tag_counts in summary_statistics.items()}

    return formatted_summary_statistics


def format_results_by_file(search_id):
    # Query all results by search_id
    all_results = ExtractedData.objects.filter(search_id=search_id)

    # Initialize a nested dictionary to hold results
    formatted_results = defaultdict(lambda: {'exact_matches': [], 'partial_matches': []})

    # Iterate through the results and group them
    for result in all_results:
        # Determine the match type and append the result to the correct category
        if result.match_type == 'exact':
            formatted_results[result.source_file_name]['exact_matches'].append(result)
        elif result.match_type == 'partial':
            formatted_results[result.source_file_name]['partial_matches'].append(result)

    return formatted_results


def serialize_formatted_results(formatted_results):
    # Initialize a dictionary to hold serialized results
    serialized_results = {}

    # Serialize each file's results
    for file_name, matches in formatted_results.items():
        serialized_results[file_name] = {
            'exact_matches': ExtractedDataSerializer(matches['exact_matches'], many=True).data,
            'partial_matches': ExtractedDataSerializer(matches['partial_matches'], many=True).data,
        }

    return serialized_results



def clean_text(text):
    """
    Clean text to remove any characters that are not XML-compatible.
    """
    if text is None:
        return ''
    # Remove NULL bytes and other control characters
    cleaned_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    return cleaned_text


def exportAsWord_using_Search_id(search_id,user_name):
            search_results = format_results_by_file(search_id)
            serialized_results = serialize_formatted_results(search_results)
            datetime_string = next(iter(serialized_results.items()))[1]["exact_matches"][0]["date_of_search"]
            datetime_string_file= datetime_string.split('.')[0]
            datetime_obj = datetime.strptime(datetime_string, "%Y-%m-%dT%H:%M:%S.%fZ")
            

            # Convert to a timezone-aware datetime object in UTC
            datetime_obj = datetime_obj.replace(tzinfo=pytz.UTC)

            word_document = export_search_results_to_word(serialized_results, search_id, datetime_obj, user_name)
            return word_document,user_name,datetime_string_file